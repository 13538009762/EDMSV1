"""Export TipTap/ProseMirror JSON to DOCX and PDF."""
from __future__ import annotations

import json
import os
import re
import logging
from html import escape as html_escape
from io import BytesIO
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from flask import current_app

from reportlab.pdfbase import pdfmetrics
from reportlab.platypus.flowables import Spacer
from reportlab.platypus.frames import Frame
from xhtml2pdf.builders.watermarks import WaterMarks
from xhtml2pdf.context import pisaContext
from xhtml2pdf.default import DEFAULT_CSS
from xhtml2pdf.document import get_encrypt_instance, pisaStory
from xhtml2pdf.files import cleanFiles
from xhtml2pdf.util import getBox
from xhtml2pdf.xhtml2pdf_reportlab import PmlBaseDoc, PmlPageTemplate
from reportlab.pdfbase.ttfonts import TTFont, TTFontFace
from reportlab.lib.fonts import addMapping

logger = logging.getLogger(__name__)
_pdf_font_cache: tuple[str, str] | None = None


def _reportlab_can_embed_font_file(path: str) -> bool:
    """
    探测文件是否能被 ReportLab 嵌入 PDF。
    Debian/Ubuntu 自带的 NotoSansCJK-Regular.ttc 常为 CFF 轮廓，会报：
    postscript outlines are not supported —— 须跳过并改用可嵌入的 TrueType（如 NotoSansSC-VF.ttf）。
    """
    if not path or not os.path.isfile(path):
        return False
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".ttc":
            TTFontFace(path, validate=0, subfontIndex=0)
        else:
            TTFontFace(path, validate=0, subfontIndex=0)
        return True
    except Exception as exc:
        logger.info("ReportLab 无法使用该字体文件，将尝试下一候选: %s (%s)", path, exc)
        return False


def _register_ttface(name: str, path: str) -> None:
    """注册 TTF/OTF/TTC；集合字体需 subfontIndex。"""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".ttc":
        pdfmetrics.registerFont(TTFont(name, path, subfontIndex=0))
    else:
        pdfmetrics.registerFont(TTFont(name, path))


def _is_arial_unicode_path(path: str) -> bool:
    base = os.path.basename(path).lower()
    return "arialuni" in base


def _get_static_dir() -> str:
    """
    统一使用 app/static。无应用上下文时也解析到同一路径（export_service 位于 app/services/）。
    """
    try:
        if current_app:
            return os.path.join(current_app.root_path, "static")
    except RuntimeError:
        pass
    services_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.normpath(os.path.join(services_dir, "..", "static"))


def get_pdf_font_spec() -> tuple[str, str]:
    """
    注册 PDF 用 Unicode 字体，返回 (CSS font-family 列表串, -pdf-font-name 主字体名)。

    此前返回 ``"PDFCJK, PDFLatin"`` 并被写成 font-family: "PDFCJK, PDFLatin"，
    会被解析成「一个名字叫 PDFCJK, PDFLatin 的字体」，ReportLab 无法匹配，回退 Helvetica 导致中文/俄语变方块。

    不在此处发起网络下载；字体须预先放入 app/static/fonts/（见 fonts-readme.txt，或运行 scripts/fetch_pdf_fonts.py）。
    """
    global _pdf_font_cache
    if _pdf_font_cache is not None:
        return _pdf_font_cache

    static_dir = _get_static_dir()
    font_dir = os.path.join(static_dir, "fonts")
    os.makedirs(font_dir, exist_ok=True)
    windir = os.environ.get("WINDIR", "C:\\Windows")

    # ---------- 正文：Arial Unicode 或本地 VF TTF（见 app/static/fonts/fonts-readme.txt）----------
    cjk_candidates = [
        os.path.join(windir, "Fonts", "arialuni.ttf"),
        os.path.join(font_dir, "NotoSansSC-VF.ttf"),
        os.path.join(font_dir, "NotoSansSC-Regular.ttf"),
        os.path.join(font_dir, "simhei.ttf"),
        os.path.join(windir, "Fonts", "simhei.ttf"),
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansSC-Regular.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf",
        # 部分镜像中的 .ttc 为 CFF，ReportLab 不支持；保留在列表末尾，由探测跳过无效文件
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    cjk_path = None
    for path in cjk_candidates:
        if not os.path.exists(path):
            continue
        if not _reportlab_can_embed_font_file(path):
            continue
        cjk_path = path
        logger.info("选用中文字体文件: %s", cjk_path)
        break
    if not cjk_path:
        raise RuntimeError(
            "未找到可用于 PDF 的中文字体（须为 ReportLab 可嵌入的 TrueType）。请将 "
            "NotoSansSC-VF.ttf 放到 app/static/fonts/（见 fonts-readme.txt），或在 backend "
            "目录运行: python scripts/fetch_pdf_fonts.py"
        )

    # 注册主字体族 PDFCJK（名称保留，便于与现有 HTML 逻辑一致）
    try:
        _register_ttface("PDFCJK", cjk_path)
        _register_ttface("PDFCJK-Bold", cjk_path)
        _register_ttface("PDFCJK-Italic", cjk_path)
        _register_ttface("PDFCJK-BoldItalic", cjk_path)
        addMapping("PDFCJK", 0, 0, "PDFCJK")
        addMapping("PDFCJK", 1, 0, "PDFCJK-Bold")
        addMapping("PDFCJK", 0, 1, "PDFCJK-Italic")
        addMapping("PDFCJK", 1, 1, "PDFCJK-BoldItalic")
        logger.info("主 PDF 字体 %s 注册成功，族名: PDFCJK", cjk_path)
    except Exception as e:
        logger.exception("主 PDF 字体注册失败 (%s): %s", cjk_path, e)
        raise RuntimeError(f"主 PDF 字体注册失败 ({cjk_path}): {e}") from e

    css_names: list[str] = ["PDFCJK"]
    pdf_name = "PDFCJK"

    # Arial Unicode 已覆盖西里尔文，无需第二套字体
    if not _is_arial_unicode_path(cjk_path):
        # ---------- 西里尔补充：SimHei/部分 CJK 字体不含俄语字形 ----------
        latin_candidates = [
            os.path.join(windir, "Fonts", "arialuni.ttf"),
            os.path.join(font_dir, "NotoSans-VF.ttf"),
            os.path.join(font_dir, "NotoSans-Regular.ttf"),
            os.path.join(font_dir, "DejaVuSans.ttf"),
        ]
        latin_path = None
        for path in latin_candidates:
            if not os.path.exists(path):
                continue
            if not _reportlab_can_embed_font_file(path):
                continue
            latin_path = path
            logger.info("找到西里尔补充字体: %s", latin_path)
            break
        if latin_path and os.path.abspath(latin_path) != os.path.abspath(cjk_path):
            try:
                _register_ttface("PDFLatin", latin_path)
                _register_ttface("PDFLatin-Bold", latin_path)
                _register_ttface("PDFLatin-Italic", latin_path)
                _register_ttface("PDFLatin-BoldItalic", latin_path)
                addMapping("PDFLatin", 0, 0, "PDFLatin")
                addMapping("PDFLatin", 1, 0, "PDFLatin-Bold")
                addMapping("PDFLatin", 0, 1, "PDFLatin-Italic")
                addMapping("PDFLatin", 1, 1, "PDFLatin-BoldItalic")
                css_names.append("PDFLatin")
                logger.info("西里尔补充字体注册成功，族名: PDFLatin")
            except Exception as e:
                logger.exception("西里尔字体注册失败 (%s): %s", latin_path, e)
                raise RuntimeError(f"西里尔字体注册失败 ({latin_path}): {e}") from e
        elif not latin_path:
            logger.warning(
                "未找到独立西里尔字体且与正文字体相同，俄语等字符可能仍显示为方块。"
            )

    css_family = ", ".join(css_names) + ", sans-serif"
    _pdf_font_cache = (css_family, pdf_name)
    return _pdf_font_cache


def get_cjk_font_family() -> str:
    """兼容旧调用：仅返回 CSS font-family 取值（含 fallback）。"""
    return get_pdf_font_spec()[0]


def _walk_text(node: dict[str, Any], parts: list[str]) -> None:
    ntype = node.get("type")
    if ntype == "text":
        t = node.get("text") or ""
        marks = node.get("marks") or []
        if any(m.get("type") == "bold" for m in marks):
            t = f"<b>{html_escape(t)}</b>"
        else:
            t = html_escape(t)
        if any(m.get("type") == "italic" for m in marks):
            t = f"<i>{t}</i>"
        if any(m.get("type") == "underline" for m in marks):
            t = f"<u>{t}</u>"
        parts.append(t)
        return
    for ch in node.get("content") or []:
        _walk_text(ch, parts)


def tiptap_json_to_plain(doc_json: str) -> str:
    try:
        root = json.loads(doc_json)
    except json.JSONDecodeError:
        return ""
    parts: list[str] = []

    def block(node: dict[str, Any]) -> None:
        t = node.get("type")
        if t == "paragraph":
            line: list[str] = []
            for ch in node.get("content") or []:
                _walk_text(ch, line)
            parts.append("".join(line))
        elif t in ("heading",):
            line = []
            for ch in node.get("content") or []:
                _walk_text(ch, line)
            lvl = (node.get("attrs") or {}).get("level") or 1
            parts.append(f"{'#' * int(lvl)} " + "".join(line))
        elif t in ("bulletList", "orderedList"):
            for item in node.get("content") or []:
                if item.get("type") == "listItem":
                    sub: list[str] = []
                    for ch in item.get("content") or []:
                        if ch.get("type") == "paragraph":
                            for c2 in ch.get("content") or []:
                                _walk_text(c2, sub)
                    parts.append("- " + "".join(sub))
        elif t == "doc":
            for ch in node.get("content") or []:
                block(ch)
        else:
            for ch in node.get("content") or []:
                if isinstance(ch, dict):
                    block(ch)

    block(root)
    return "\n".join(parts)


def tiptap_json_to_html(
    doc_json: str,
    font_family_css: str = "PDFCJK, sans-serif",
    pdf_font_name: str = "PDFCJK",
    page_settings: dict = None,
) -> str:
    """将 TipTap JSON 转换为 HTML；font_family_css 为合法 CSS 列表（勿把整个列表包在一对引号里）。"""
    try:
        root = json.loads(doc_json)
    except json.JSONDecodeError:
        return "<p></p>"

    # Default page settings
    ps = {
        "orientation": "portrait",
        "marginTop": 40,
        "marginBottom": 40,
        "paperFormat": "A4",
        "showPageNumber": True
    }
    if page_settings:
        ps.update(page_settings)

    # Map margin names to values
    margin_top = f"{ps.get('marginTop', 40)}mm"
    margin_bottom = f"{ps.get('marginBottom', 40)}mm"
    size_val = f"{ps.get('paperFormat', 'A4')} {ps.get('orientation', 'portrait')}"

    def inline(node: dict[str, Any]) -> str:
        if node.get("type") == "text":
            t = html_escape(node.get("text") or "")
            for m in node.get("marks") or []:
                mt = m.get("type")
                if mt == "bold":
                    t = f"<strong>{t}</strong>"
                elif mt == "italic":
                    t = f"<em>{t}</em>"
                elif mt == "underline":
                    t = f"<u>{t}</u>"
                elif mt == "strike":
                    t = f"<s>{t}</s>"
                elif mt == "textStyle":
                    attrs = m.get("attrs") or {}
                    color = attrs.get("color")
                    if color:
                        t = f'<span style="color:{html_escape(color)}">{t}</span>'
            return t
        return ""

    def block(node: dict[str, Any]) -> str:
        t = node.get("type")
        if t == "paragraph":
            inner = "".join(
                inline(c) if c.get("type") == "text" else "" for c in node.get("content") or []
            )
            style = ""
            ta = (node.get("attrs") or {}).get("textAlign")
            if ta:
                style = f' style="text-align:{html_escape(ta)}"'
            return f"<p{style}>{inner or '&nbsp;'}</p>"
        if t == "heading":
            lvl = int((node.get("attrs") or {}).get("level") or 1)
            inner = "".join(
                inline(c) if c.get("type") == "text" else "" for c in node.get("content") or []
            )
            return f"<h{lvl}>{inner}</h{lvl}>"
        if t == "bulletList":
            items = "".join(block(c) for c in node.get("content") or [])
            return f"<ul>{items}</ul>"
        if t == "orderedList":
            items = "".join(block(c) for c in node.get("content") or [])
            return f"<ol>{items}</ol>"
        if t == "listItem":
            inner = "".join(block(c) for c in node.get("content") or [])
            return f"<li>{inner}</li>"
        if t == "table":
            inner = "".join(block(c) for c in node.get("content") or [])
            return f"<table border='1' cellpadding='6' cellspacing='0' style='border-collapse: collapse; width: 100%;'>{inner}</table>"
        if t == "tableRow":
            inner = "".join(block(c) for c in node.get("content") or [])
            return f"<tr>{inner}</tr>"
        if t in ("tableCell", "tableHeader"):
            tag = "th" if t == "tableHeader" else "td"
            inner = "".join(block(c) for c in node.get("content") or [])
            colspan = (node.get("attrs") or {}).get("colspan", 1)
            rowspan = (node.get("attrs") or {}).get("rowspan", 1)
            attrs_str = f' colspan="{colspan}"' if colspan > 1 else ""
            attrs_str += f' rowspan="{rowspan}"' if rowspan > 1 else ""
            return f"<{tag}{attrs_str}>{inner}</{tag}>"
        if t == "image":
            attrs = node.get("attrs") or {}
            src = attrs.get("src") or ""
            align = attrs.get("textAlign")
            width = attrs.get("width") or "100%"
            style = f"max-width: 100%; width: {html_escape(width)};"
            wrapper_style = "display: block;"
            if align:
                wrapper_style += f" text-align:{align};"
            return f"<div style='{wrapper_style}'><img src='{html_escape(src)}' style='{style}' /></div>"
        if t == "pageBreak":
            return "<pdf:nextpage />"
        if t == "doc":
            return "".join(block(c) for c in node.get("content") or [])
        return "".join(block(c) for c in node.get("content") or [])

    body_html = block(root)

    # CSS for xhtml2pdf Page Numbering and Settings
    style = f"""
    <style>
        @page {{
            size: {size_val};
            margin-top: {margin_top};
            margin-bottom: {margin_bottom};
            margin-left: 25mm;
            margin-right: 25mm;
            @frame footer {{
                -pdf-frame-content: footer_content;
                bottom: 10mm;
                margin-left: 25mm;
                margin-right: 25mm;
                height: 10mm;
            }}
        }}
        body {{
            font-family: {font_family_css};
            -pdf-font-name: {pdf_font_name};
            -pdf-encoding: Identity-H;
            font-size: 14px;
            line-height: 1.5;
            color: #333;
            margin: 0;
            padding: 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin-top: 10px;
            margin-bottom: 10px;
        }}
        th, td {{
            border: 1px solid black;
            padding: 6px;
            word-wrap: break-word;
            font-family: {font_family_css};
        }}
        img {{
            max-width: 100%;
        }}
        #footer_content {{
            text-align: center;
            color: #555;
            font-size: 11px;
            font-family: {font_family_css};
        }}
    </style>
    """
    
    footer = ""
    show_pg = ps.get("showPageNumber")
    # Robust check for boolean or string "true"
    if show_pg is True or str(show_pg).lower() == "true":
        footer = f'<div id="footer_content">- <pdf:pagenumber /> -</div>'

    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        {style}
    </head>
    <body style="font-family: {font_family_css};">
        {footer}
        {body_html}
    </body>
    </html>
    """



def _set_run_font(run, font_name: str = "SimHei"):
    """为 docx 的 run 设置字体名称，确保支持中文和俄语"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def export_docx_bytes(doc_json: str, page_settings: dict = None) -> bytes:
    """从 TipTap JSON 生成 DOCX 文档，使用通用字体（SimHei）以支持中俄语"""
    from docx.shared import Mm
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

    data = json.loads(doc_json) if doc_json else {"type": "doc", "content": []}
    d = DocxDocument()

    # Apply page settings to DOCX
    if page_settings:
        section = d.sections[0]
        
        # 1. 设置纸张方向
        if page_settings.get("orientation") == "landscape":
            section.orientation = 1 # landscape
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
        
        # 2. 设置上下边距 (读取前端传来的毫米数值)
        if "marginTop" in page_settings:
            section.top_margin = Mm(page_settings["marginTop"])
        if "marginBottom" in page_settings:
            section.bottom_margin = Mm(page_settings["marginBottom"])
            
        # 3. 设置左右边距 (前端如果没传，默认固定 25 毫米)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)

        # 4. 强行注入 Word 底层动态页码
        show_pg = page_settings.get("showPageNumber")
        if show_pg is True or str(show_pg).lower() == "true":
            def add_page_number(run):
                # 构造 Word 底层 XML 标签用于显示自增页码 "PAGE"
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = "PAGE"
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'separate')
                fldChar3 = OxmlElement('w:fldChar')
                fldChar3.set(qn('w:fldCharType'), 'end')
                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)
                run._r.append(fldChar3)

            # 遍历并写入到每一节(Section)的页脚中
            for sec in d.sections:
                footer = sec.footer
                # 如果原有页脚段落存在则取第一个，否则新建
                p = footer.paragraphs[0] if len(footer.paragraphs) > 0 else footer.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.clear()  # 清除垃圾空文本防冲突
                
                run = p.add_run()
                # 给页码设定中文字体，防乱码
                run.font.name = "SimHei"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "SimHei")
                
                # 写入动态页码
                add_page_number(run)

    def process_marks(run, marks):
        for m in marks or []:
            mt = m.get("type")
            if mt == "bold":
                run.bold = True
            elif mt == "italic":
                run.italic = True
            elif mt == "underline":
                run.underline = True
            elif mt == "strike":
                run.font.strike = True
            elif mt == "textStyle":
                col = (m.get("attrs") or {}).get("color")
                if col and col.startswith("#") and len(col) >= 7:
                    try:
                        r = int(col[1:3], 16)
                        g = int(col[3:5], 16)
                        b = int(col[5:7], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
                    except ValueError:
                        pass

    def add_block_from_node(node: dict[str, Any], container) -> None:
        t = node.get("type")
        if t == "paragraph":
            p = container.add_paragraph()
            for ch in node.get("content") or []:
                if ch.get("type") == "text":
                    run = p.add_run(ch.get("text") or "")
                    process_marks(run, ch.get("marks"))
                    _set_run_font(run, "SimHei")
            ta = (node.get("attrs") or {}).get("textAlign")
            if ta == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif ta == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        elif t == "heading":
            lvl = int((node.get("attrs") or {}).get("level") or 1)
            try:
                p = container.add_heading(level=min(lvl, 3))
            except AttributeError:
                p = container.add_paragraph()
                r = p.add_run(f"{'#' * lvl} ")
                _set_run_font(r, "SimHei")
                r.bold = True

            for ch in node.get("content") or []:
                if ch.get("type") == "text":
                    run = p.add_run(ch.get("text") or "")
                    process_marks(run, ch.get("marks"))
                    _set_run_font(run, "SimHei")

        elif t in ("bulletList", "orderedList"):
            style = "List Bullet" if t == "bulletList" else "List Number"
            for item in node.get("content") or []:
                if item.get("type") != "listItem":
                    continue
                for sub in item.get("content") or []:
                    if sub.get("type") == "paragraph":
                        p = container.add_paragraph(style=style)
                        for ch in sub.get("content") or []:
                            if ch.get("type") == "text":
                                run = p.add_run(ch.get("text") or "")
                                process_marks(run, ch.get("marks"))
                                _set_run_font(run, "SimHei")

        elif t == "table":
            row_nodes = [r for r in node.get("content", []) if r.get("type") == "tableRow"]
            if not row_nodes:
                return

            first_row_cells = [c for c in row_nodes[0].get("content", []) if c.get("type") in ("tableCell", "tableHeader")]
            max_cols = sum((c.get("attrs") or {}).get("colspan", 1) for c in first_row_cells)
            if max_cols == 0:
                return

            table = container.add_table(rows=len(row_nodes), cols=max_cols)
            table.style = 'Table Grid'

            for r_idx, r_node in enumerate(row_nodes):
                c_idx = 0
                cell_nodes = [c for c in r_node.get("content", []) if c.get("type") in ("tableCell", "tableHeader")]
                for cell_node in cell_nodes:
                    if c_idx >= max_cols:
                        break

                    colspan = (cell_node.get("attrs") or {}).get("colspan", 1)
                    if colspan > 1:
                        end_idx = min(c_idx + colspan - 1, max_cols - 1)
                        table.cell(r_idx, c_idx).merge(table.cell(r_idx, end_idx))

                    cell = table.cell(r_idx, c_idx)
                    if len(cell.paragraphs) > 0:
                        cell.paragraphs[0].text = ""

                    first_p = True
                    for ch in cell_node.get("content", []):
                        if ch.get("type") == "paragraph":
                            if first_p and len(cell.paragraphs) > 0:
                                p = cell.paragraphs[0]
                                first_p = False
                            else:
                                p = cell.add_paragraph()

                            for text_node in ch.get("content", []):
                                if text_node.get("type") == "text":
                                    run = p.add_run(text_node.get("text") or "")
                                    process_marks(run, text_node.get("marks"))
                                    _set_run_font(run, "SimHei")

                            ta = (ch.get("attrs") or {}).get("textAlign")
                            if ta == "center":
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif ta == "right":
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    c_idx += colspan

        elif t == "pageBreak":
            try:
                container.add_page_break()
            except AttributeError:
                container.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        elif t == "image":
            attrs = node.get("attrs") or {}
            src = attrs.get("src")
            align = attrs.get("textAlign")
            width_str = attrs.get("width")
            
            if src:
                try:
                    from docx.shared import Inches, Mm
                    
                    # 🖼️ Determine local path
                    img_path = None
                    if src.startswith('/static/'):
                         static_dir = _get_static_dir()
                         # src like "/static/images/..." 
                         # static_dir like ".../static"
                         # Extract relative path: src[8:] is "images/..."
                         img_path = os.path.join(static_dir, src[8:])
                    
                    if img_path and os.path.exists(img_path):
                        p = container.add_paragraph()
                        # Handle alignment
                        if align == "center":
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif align == "right":
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        
                        run = p.add_run()
                        # Handle width
                        w_val = None
                        section = d.sections[0]
                        max_w = section.page_width - section.left_margin - section.right_margin
                        
                        if width_str:
                            if width_str.endswith("%"):
                                percent = float(width_str[:-1]) / 100.0
                                w_val = max_w * percent
                            elif width_str.endswith("px"):
                                w_val = Mm(float(width_str[:-2]) * 0.264)
                        
                        # Apply to docx
                        if w_val:
                            run.add_picture(img_path, width=min(w_val, max_w))
                        else:
                            run.add_picture(img_path, width=min(Inches(5.0), max_w))
                except Exception as e:
                    logger.error(f"DOCX image injection failed: {e}")

        elif t == "doc":

            for ch in node.get("content") or []:
                add_block_from_node(ch, container)
        else:
            for ch in node.get("content") or []:
                if isinstance(ch, dict):
                    add_block_from_node(ch, container)

    add_block_from_node(data, d)
    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


def _fetch_resources(uri, rel):
    if uri.startswith('/static/'):
        static_dir = _get_static_dir()
        return os.path.join(static_dir, uri[8:])
    return uri


def _write_pdf_from_html(html: str, dest: BytesIO) -> None:
    """
    生成 PDF。xhtml2pdf 只认 pisaContext.fontList，不能仅靠 pdfmetrics.registerFont，
    否则 font-family 会回退 Helvetica，中文/俄语成方块。
    """
    import io

    base_path = _get_static_dir()
    context = pisaContext(base_path, debug=0, capacity=100 * 1024)
    context.pathCallback = _fetch_resources

    for reg_name in ("PDFCJK", "PDFLatin"):
        if reg_name in pdfmetrics.getRegisteredFontNames():
            context.registerFont(reg_name)

    context = pisaStory(
        html,
        path=base_path,
        link_callback=_fetch_resources,
        debug=0,
        default_css=DEFAULT_CSS,
        xhtml=False,
        encoding="utf-8",
        context=context,
        xml_output=None,
    )

    if context.err:
        for mode, line, msg, _record in context.log:
            if mode == "error":
                logger.error("xhtml2pdf line %s: %s", line, msg)

    if not context.story:
        context.story = [Spacer(1, 1)]

    for frag, anchor in context.anchorFrag:
        if anchor not in context.anchorName:
            frag.link = None

    buf = io.BytesIO()
    doc = PmlBaseDoc(
        buf,
        pagesize=context.pageSize,
        author=context.meta["author"].strip(),
        subject=context.meta["subject"].strip(),
        keywords=[
            x.strip()
            for x in (context.meta.get("keywords") or "").strip().split(",")
            if x.strip()
        ],
        title=context.meta["title"].strip(),
        showBoundary=0,
        encrypt=get_encrypt_instance(None),
        allowSplitting=1,
    )

    if "body" in context.templateList:
        body = context.templateList["body"]
        del context.templateList["body"]
    else:
        x, y, w, h = getBox("1cm 1cm -1cm -1cm", context.pageSize)
        body = PmlPageTemplate(
            id="body",
            frames=[
                Frame(
                    x,
                    y,
                    w,
                    h,
                    id="body",
                    leftPadding=0,
                    rightPadding=0,
                    bottomPadding=0,
                    topPadding=0,
                )
            ],
            pagesize=context.pageSize,
        )

    doc.addPageTemplates([body] + list(context.templateList.values()))

    if context.multiBuild:
        doc.multiBuild(context.story)
    else:
        doc.build(context.story)

    output = io.BytesIO()
    output, has_bg = WaterMarks.process_doc(context, buf, output)
    if not has_bg:
        output = buf

    data = output.getvalue()
    dest.write(data)
    cleanFiles()


def export_pdf_bytes(doc_json: str, page_settings: dict = None) -> bytes:
    """从 TipTap JSON 生成 PDF 文档，使用已注册的通用字体"""
    font_family_css, pdf_font_name = get_pdf_font_spec()
    html = tiptap_json_to_html(
        doc_json or "{}",
        font_family_css=font_family_css,
        pdf_font_name=pdf_font_name,
        page_settings=page_settings,
    )
    out = BytesIO()
    _write_pdf_from_html(html, out)
    return out.getvalue()


def apply_punctuation_fixes(text: str) -> str:
    text = re.sub(r"，{2,}", "，", text)
    text = re.sub(r"。{2,}", "。", text)
    text = text.replace("。。", "。")
    return text